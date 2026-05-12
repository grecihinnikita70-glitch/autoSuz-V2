from pathlib import Path

from docx import Document as DocxDocument

from app.services import pdf_parser
from app.services.docx_parser import parse_docx
from app.services.extraction_types import ExtractedField, ExtractionResult
from app.services.pdf_parser import SCAN_WARNING, parse_pdf
from app.services.regex_extractor import (
    extract_amounts,
    extract_all_fields,
    extract_contract_date,
    extract_contract_number,
    extract_insurance_period,
    extract_insurance_fields,
    extract_pledge_agreements,
)
from app.services.text_normalizer import (
    normalize_dates_text,
    normalize_quotes,
    normalize_spaces,
    normalize_text,
)
from app.services.validator import has_allowed_extension


class FakePdf:
    def __init__(self, pages):
        self.pages = pages

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_value, traceback):
        return False


class FakePdfPage:
    def __init__(self, text, words):
        self._text = text
        self._words = words

    def extract_text(self):
        return self._text

    def extract_words(self):
        return self._words


def test_has_allowed_extension_accepts_supported_files():
    assert has_allowed_extension("contract.docx")
    assert has_allowed_extension("CONTRACT.DOCX")
    assert has_allowed_extension("contract.pdf")


def test_has_allowed_extension_rejects_unknown_format():
    assert not has_allowed_extension("contract.txt")


def test_normalize_text_compacts_spaces_and_empty_lines():
    raw_text = "  Insurance   contract\r\n\r\n\r\n  Number\t123  "

    assert normalize_text(raw_text) == "Insurance contract\n\nNumber 123"


def test_normalize_spaces_keeps_paragraphs_and_number_sign():
    raw_text = "  Policy\u00a0\u00a0№\t123  \n\n\n  Holder:   Ivanov  "

    assert normalize_spaces(raw_text) == "Policy № 123\n\nHolder: Ivanov"


def test_normalize_quotes_uses_regular_double_quotes():
    raw_text = "\u00abAlpha\u00bb, \u201cBeta\u201d, \"Gamma\""

    assert normalize_quotes(raw_text) == '"Alpha", "Beta", "Gamma"'


def test_normalize_dates_text_formats_simple_numeric_dates():
    raw_text = "Dates: 1 / 2 / 2025, 03-04-25, 05.06.2024"

    assert normalize_dates_text(raw_text) == (
        "Dates: 01.02.2025, 03.04.25, 05.06.2024"
    )


def test_normalize_text_does_not_break_contract_numbers():
    numbers = [
        "SGZPG-80070101601030VTB",
        "СВЛ/111111-111111-303",
        "До38-ЦВ-730710/2020/00125",
        "50912/444/00036/25",
        "441-581-246458/25",
    ]
    raw_text = "\n".join(numbers)

    assert normalize_text(raw_text) == raw_text


def test_regex_extractor_is_disabled_for_first_mvp():
    assert extract_insurance_fields("Any text") == {}


def test_extract_contract_number_from_real_examples():
    examples = [
        ("ДОГОВОР № SGZPG-80070101601030VTB", "SGZPG-80070101601030VTB"),
        ("ДОГОВОР № SGZPG-01260000004241VTB", "SGZPG-01260000004241VTB"),
        ("ПОЛИС № 555 РТ 1234 SMBRE", "555 РТ 1234 SMBRE"),
        ("ПОЛИС № SGZPG-06260000004097SMBRE", "SGZPG-06260000004097SMBRE"),
        (
            "ДОГОВОР СТРАХОВАНИЯ ИМУЩЕСТВА № 78542/919/50117/25",
            "78542/919/50117/25",
        ),
        (
            "ПОЛИС по страхованию имущества № 441-581-246458/25",
            "441-581-246458/25",
        ),
    ]

    for text, expected_value in examples:
        extracted_field = extract_contract_number(text)

        assert extracted_field.value == expected_value
        assert extracted_field.confidence == 0.95
        assert extracted_field.method == "regex"
        assert extracted_field.source_text is not None
        assert expected_value in extracted_field.source_text


def test_extract_contract_number_fallback_after_number_sign():
    extracted_field = extract_contract_number("№ 252500-141-000173")

    assert extracted_field.value == "252500-141-000173"
    assert extracted_field.confidence == 0.75
    assert extracted_field.method == "regex"


def test_extract_contract_number_trims_text_after_number():
    examples = [
        ("ИМУЩЕСТВА ЮРИДИЧЕСКИХ ЛИЦ № 252500-141-000173 г. Краснодар", "252500-141-000173"),
        (
            "по страхованию имущества № 441-581-246458/25 Следующие объекты и условия",
            "441-581-246458/25",
        ),
        ('"АЛЬФАКАЛЬКУЛЯТОР ИЮЛ" № 50912/444/00036/25 32858734', "50912/444/00036/25"),
    ]

    for text, expected_value in examples:
        assert extract_contract_number(text).value == expected_value


def test_extract_contract_number_ignores_credit_and_pledge_contracts():
    text = """
    Кредитный договор № 111-222-333
    Договор залога № 444-555-666
    ДОГОВОР СТРАХОВАНИЯ ИМУЩЕСТВА № 78542/919/50117/25
    """

    extracted_field = extract_contract_number(text)

    assert extracted_field.value == "78542/919/50117/25"
    assert "СТРАХОВАНИЯ" in extracted_field.source_text


def test_extract_contract_number_returns_not_found_for_only_credit_contract():
    text = """
    Кредитный договор № 111-222-333
    Договор залога № 444-555-666
    """

    extracted_field = extract_contract_number(text)

    assert extracted_field.value is None
    assert extracted_field.confidence == 0.0
    assert extracted_field.method == "not_found"


def test_extract_contract_number_ignores_appendix_and_table_markers():
    text = """
    Приложение № 1 к Договору
    Уд № тек. Состояния 2972912847
    """

    extracted_field = extract_contract_number(text)

    assert extracted_field.value is None
    assert extracted_field.method == "not_found"


def test_extract_contract_number_searches_only_document_start():
    text = "A" * 3100 + "\nДОГОВОР № SGZPG-80070101601030VTB"

    extracted_field = extract_contract_number(text)

    assert extracted_field.value is None
    assert extracted_field.method == "not_found"


def test_extract_contract_date_from_supported_formats():
    examples = [
        ("г. Волгоград «13» марта 2026 г.", "13.03.2026"),
        ("г. Астрахань «25» ноября 2025 г.", "25.11.2025"),
        ("г. Санкт-Петербург, 05 марта 2026 г.", "05.03.2026"),
        ("«11» сентября 2025 г.", "11.09.2025"),
        ("«10» декабря 2025 г.", "10.12.2025"),
        ("06 ноября 2025 года", "06.11.2025"),
    ]

    for text, expected_value in examples:
        extracted_field = extract_contract_date(text)

        assert extracted_field.value == expected_value
        assert extracted_field.method == "regex"
        assert extracted_field.confidence == 0.9
        assert extracted_field.source_text is not None
        assert text.split()[0] in extracted_field.source_text


def test_extract_contract_date_returns_not_found():
    extracted_field = extract_contract_date("Документ без даты договора")

    assert extracted_field.value is None
    assert extracted_field.confidence == 0.0
    assert extracted_field.method == "not_found"


def test_extract_contract_date_skips_power_of_attorney_date():
    text = """
    Доверенность от «01» января 2026 г.
    ДОГОВОР СТРАХОВАНИЯ
    г. Волгоград «13» марта 2026 г.
    """

    extracted_field = extract_contract_date(text)

    assert extracted_field.value == "13.03.2026"
    assert "Волгоград" in extracted_field.source_text


def test_extract_contract_date_skips_insurance_rules_date():
    text = """
    Правила страхования от «01» января 2020 г.
    ПОЛИС № SGZPG-06260000004097SMBRE
    «10» декабря 2025 г.
    """

    extracted_field = extract_contract_date(text)

    assert extracted_field.value == "10.12.2025"
    assert "декабря" in extracted_field.source_text


def test_extract_contract_date_searches_only_document_start():
    text = "A" * 4100 + "\n«10» декабря 2025 г."

    extracted_field = extract_contract_date(text)

    assert extracted_field.value is None
    assert extracted_field.method == "not_found"


def test_extract_insurance_period_from_policy_validity_dates():
    text = (
        "Срок действия полиса с 29.03.2026, 00 ч. 00 мин "
        "по 28.03.2027, 24 ч. 00 мин"
    )

    period_start, period_end = extract_insurance_period(text)

    assert period_start.value == "29.03.2026"
    assert period_end.value == "28.03.2027"
    assert period_start.confidence == 0.95
    assert period_end.confidence == 0.95
    assert period_start.method == "regex"
    assert period_start.source_text == period_end.source_text


def test_extract_insurance_period_from_explicit_section():
    text = (
        "ПЕРИОД СТРАХОВАНИЯ\n"
        "С 00:00 часов «03» декабря 2025 г. "
        "по 24:00 часов «02» декабря 2026 г."
    )

    period_start, period_end = extract_insurance_period(text)

    assert period_start.value == "03.12.2025"
    assert period_end.value == "02.12.2026"
    assert period_start.confidence == 0.95
    assert "ПЕРИОД СТРАХОВАНИЯ" in period_start.source_text


def test_extract_insurance_period_from_general_range_template():
    text = 'с «30» ноября 2025 года по «29» ноября 2026 года'

    period_start, period_end = extract_insurance_period(text)

    assert period_start.value == "30.11.2025"
    assert period_end.value == "29.11.2026"
    assert period_start.confidence == 0.75
    assert period_end.confidence == 0.75


def test_extract_insurance_period_supports_text_dates_with_times():
    text = (
        "С 00:00 часов «19» октября 2025 г. "
        "по 24:00 часов «18» октября 2026 г."
    )

    period_start, period_end = extract_insurance_period(text)

    assert period_start.value == "19.10.2025"
    assert period_end.value == "18.10.2026"
    assert period_start.confidence == 0.75


def test_extract_insurance_period_ignores_credit_pledge_and_power_of_attorney():
    text = """
    Кредитный договор действует с 01.01.2025 по 01.01.2026.
    Договор залога действует с 02.01.2025 по 02.01.2026.
    Доверенность действует с 03.01.2025 по 03.01.2026.
    ПЕРИОД СТРАХОВАНИЯ
    с «30» ноября 2025 года по «29» ноября 2026 года
    """

    period_start, period_end = extract_insurance_period(text)

    assert period_start.value == "30.11.2025"
    assert period_end.value == "29.11.2026"
    assert period_start.confidence == 0.95


def test_extract_insurance_period_returns_not_found():
    period_start, period_end = extract_insurance_period("Документ без периода")

    assert period_start.value is None
    assert period_start.method == "not_found"
    assert period_start.confidence == 0.0
    assert period_end.value is None
    assert period_end.method == "not_found"
    assert period_end.confidence == 0.0


def test_extract_amounts_from_explicit_total_phrases():
    text = """
    Итого общая страховая сумма по Договору: 73 353 000,00
    Итого общая страховая премия по настоящему Договору: 73 353,00
    """

    amounts = extract_amounts(text)

    assert amounts["total_insured_amount"].value == "73353000.00"
    assert amounts["total_insurance_premium"].value == "73353.00"
    assert amounts["total_insured_amount"].confidence == 0.95
    assert amounts["total_insurance_premium"].confidence == 0.95
    assert amounts["total_insured_amount"].method == "regex"


def test_extract_amounts_from_total_table_line():
    text = "Итого: 37 037 600,00    40 653,18"

    amounts = extract_amounts(text)

    assert amounts["total_insured_amount"].value == "37037600.00"
    assert amounts["total_insurance_premium"].value == "40653.18"
    assert amounts["total_insured_amount"].confidence == 0.75
    assert amounts["total_insurance_premium"].confidence == 0.75


def test_extract_amounts_prefers_aggregate_premium_over_installment():
    text = """
    Первый страховой взнос: 12 000,00
    Совокупная страховая премия за весь срок действия Полиса: 114 436,02
    Страховая сумма по договору: 127 151 134,02
    """

    amounts = extract_amounts(text)

    assert amounts["total_insurance_premium"].value == "114436.02"
    assert amounts["total_insurance_premium"].source_text.startswith("Совокупная")
    assert amounts["total_insured_amount"].value == "127151134.02"


def test_extract_amounts_from_soft_insured_amount_and_premium_phrases():
    text = """
    Страховая сумма по объектам страхования составляет 127 151 134,02
    Страховая премия по Полису составляет 114 436,02
    """

    amounts = extract_amounts(text)

    assert amounts["total_insured_amount"].value == "127151134.02"
    assert amounts["total_insurance_premium"].value == "114436.02"
    assert amounts["total_insured_amount"].confidence == 0.85
    assert amounts["total_insurance_premium"].confidence == 0.85


def test_extract_amounts_does_not_use_franchise_as_insured_amount():
    text = """
    Франшиза по договору составляет 900 000,00
    Страховая премия по договору составляет 10 000,00
    """

    amounts = extract_amounts(text)

    assert amounts["total_insured_amount"].value is None
    assert amounts["total_insured_amount"].method == "not_found"
    assert amounts["total_insurance_premium"].value == "10000.00"


def test_extract_amounts_does_not_use_installment_as_premium():
    text = """
    Страховой взнос по графику: 20 000,00
    Страховая сумма по договору: 1 000 000,00
    """

    amounts = extract_amounts(text)

    assert amounts["total_insurance_premium"].value is None
    assert amounts["total_insurance_premium"].method == "not_found"
    assert amounts["total_insured_amount"].value == "1000000.00"


def test_extract_amounts_uses_largest_money_as_insured_amount_when_needed():
    text = """
    Объект 1: 500 000,00
    Объект 2: 1 500 000,00
    Страховая премия: 15 000,00
    """

    amounts = extract_amounts(text)

    assert amounts["total_insured_amount"].value == "1500000.00"
    assert amounts["total_insured_amount"].confidence == 0.6
    assert amounts["total_insurance_premium"].value == "15000.00"


def test_extract_amounts_explicit_premium_is_not_blocked_by_later_franchise():
    text = """
    Итого общая страховая сумма по Договору: 73 353 000,00
    Итого общая страховая премия по настоящему Договору: 73 353,00
    ФРАНШИЗА:
    Франшиза по настоящему Договору не установлена.
    """

    amounts = extract_amounts(text)

    assert amounts["total_insurance_premium"].value == "73353.00"
    assert amounts["total_insurance_premium"].confidence == 0.95


def test_extract_amounts_soft_premium_uses_smaller_value_in_table_row():
    text = """
    Страховая сумма по договору: 73 353 000,00
    Страховая премия по рискам 44 411 000,00 44 411,00
    """

    amounts = extract_amounts(text)

    assert amounts["total_insured_amount"].value == "73353000.00"
    assert amounts["total_insurance_premium"].value == "44411.00"


def test_extract_amounts_largest_fallback_ignores_percent_rates():
    text = "за каждый месяц 0,83%"

    amounts = extract_amounts(text)

    assert amounts["total_insured_amount"].value is None
    assert amounts["total_insured_amount"].method == "not_found"


def test_extract_pledge_agreements_from_real_examples():
    text = """
    является предметом залога по Договорам залога недвижимого имущества СВЛ/111111-111111-303 от 01.01.2025
    по Договору залога № До38-ЦВ-730710/2020/00125 от 02.01.2025
    Договору об ипотеке зданий и земельных участков № 1/26/21/10 от 03.01.2025
    Договорам о последующей ипотеке зданий и земельных участков № СНЛ/555555-559955-301 от 04.01.2025
    Договора о последующей ипотеке здания № СНЛ/555424-363937-302 от 05.01.2025
    Договора об ипотеке № КЗ/30-00/22-00066-301
    """

    fields = extract_pledge_agreements(text)

    assert [field.value for field in fields] == [
        "СВЛ/111111-111111-303",
        "До38-ЦВ-730710/2020/00125",
        "1/26/21/10",
        "СНЛ/555555-559955-301",
        "СНЛ/555424-363937-302",
        "КЗ/30-00/22-00066-301",
    ]
    assert all(field.method == "regex" for field in fields)
    assert all(field.confidence == 0.9 for field in fields)
    assert "Договора" in fields[0].source_text


def test_extract_pledge_agreements_ignores_credit_agreements():
    text = """
    кредитное соглашение № КЗ/30-00/22-00066-301
    кредитный договор № СВЛ/111111-111111-303
    """

    assert extract_pledge_agreements(text) == []


def test_extract_pledge_agreements_deduplicates_and_keeps_order():
    text = """
    Договору залога № СВЛ/111111-111111-303
    Договора об ипотеке № CHLI/555125-512534-301
    Договору залога № СВЛ/111111-111111-303 повторно указан.
    """

    fields = extract_pledge_agreements(text)

    assert [field.value for field in fields] == [
        "СВЛ/111111-111111-303",
        "CHLI/555125-512534-301",
    ]


def test_extract_pledge_agreements_prefers_full_number_over_prefix():
    text = """
    Договору залога № СВЛ/111111-111111
    Договору залога № СВЛ/111111-111111-303
    """

    fields = extract_pledge_agreements(text)

    assert [field.value for field in fields] == ["СВЛ/111111-111111-303"]


def test_extract_pledge_agreements_multiple_numbers_in_one_paragraph():
    text = (
        "Предметом залога по договорам залога СВЛ/111111-111111-303 "
        "и СНЛ/555424-363937-302 является имущество страхователя."
    )

    fields = extract_pledge_agreements(text)

    assert [field.value for field in fields] == [
        "СВЛ/111111-111111-303",
        "СНЛ/555424-363937-302",
    ]


def test_extract_all_fields_from_artificial_contract_text():
    text = """
    ДОГОВОР СТРАХОВАНИЯ ИМУЩЕСТВА № 78542/919/50117/25
    г. Волгоград «13» марта 2026 г.

    Страховщик: АО "Страховая компания", ИНН 7701234567
    Страхователь: ООО "Ромашка", ИНН 771234567890

    ПЕРИОД СТРАХОВАНИЯ
    с «30» ноября 2025 года по «29» ноября 2026 года

    Итого общая страховая сумма по Договору: 73 353 000,00
    Итого общая страховая премия по настоящему Договору: 73 353,00

    Имущество является предметом залога по Договору залога № СВЛ/111111-111111-303 от 01.01.2025.
    """

    result = extract_all_fields(text)
    result_dict = result.to_dict()
    result_json = result.to_json()

    assert isinstance(result, ExtractionResult)
    assert result.contract_number.value == "78542/919/50117/25"
    assert result.contract_date.value == "13.03.2026"
    assert result.period_start.value == "30.11.2025"
    assert result.period_end.value == "29.11.2026"
    assert result.insurer_inn.value == "7701234567"
    assert result.policyholder_inn.value == "771234567890"
    assert result.total_insured_amount.value == "73353000.00"
    assert result.total_insurance_premium.value == "73353.00"
    assert [field.value for field in result.pledge_agreements] == ["СВЛ/111111-111111-303"]
    assert result_dict["contract_number"]["value"] == "78542/919/50117/25"
    assert '"contract_number"' in result_json


def test_extract_all_fields_adds_validator_warnings():
    text = """
    ДОГОВОР № SGZPG-80070101601030VTB
    Страховщик: АО "Страховая компания", ИНН 123456789
    Страхователь: ООО "Ромашка", ИНН 1234567890123
    Срок действия полиса с 29.03.2027 по 28.03.2026
    Итого общая страховая сумма по Договору: 10 000,00
    Итого общая страховая премия по настоящему Договору: 20 000,00
    """

    result = extract_all_fields(text)

    assert "ИНН должен содержать 10 или 12 цифр." in result.insurer_inn.warnings
    assert "ИНН должен содержать 10 или 12 цифр." in result.policyholder_inn.warnings
    assert "Дата не распознана." in result.contract_date.warnings
    assert "Страховая премия больше страховой суммы." in result.total_insurance_premium.warnings
    assert "Не найден ни один договор залога." in result.contract_number.warnings
    assert "Дата окончания периода раньше даты начала." in result.period_end.warnings


def test_extracted_field_to_dict():
    extracted_field = ExtractedField(
        name="contract_number",
        label="Contract number",
        value="SGZPG-80070101601030VTB",
        confidence=0.92,
        method="regex",
        source_text="Contract № SGZPG-80070101601030VTB",
        warnings=["Check manually"],
    )

    assert extracted_field.to_dict() == {
        "name": "contract_number",
        "label": "Contract number",
        "value": "SGZPG-80070101601030VTB",
        "confidence": 0.92,
        "method": "regex",
        "source_text": "Contract № SGZPG-80070101601030VTB",
        "warnings": ["Check manually"],
    }


def test_extracted_field_defaults_to_not_found():
    extracted_field = ExtractedField(
        name="contract_date",
        label="Contract date",
    )

    assert extracted_field.value is None
    assert extracted_field.confidence == 0.0
    assert extracted_field.method == "not_found"
    assert extracted_field.warnings == []


def test_extracted_field_validates_method_and_confidence():
    try:
        ExtractedField(
            name="contract_number",
            label="Contract number",
            value="123",
            confidence=1.2,
            method="regex",
        )
    except ValueError as error:
        assert "confidence" in str(error)
    else:
        raise AssertionError("Expected ValueError for invalid confidence")

    try:
        ExtractedField(
            name="contract_number",
            label="Contract number",
            value="123",
            confidence=0.5,
            method="unknown",
        )
    except ValueError as error:
        assert "method" in str(error)
    else:
        raise AssertionError("Expected ValueError for invalid method")


def test_extraction_result_to_dict():
    def missing_field(name, label):
        return ExtractedField(name=name, label=label)

    result = ExtractionResult(
        contract_number=ExtractedField(
            name="contract_number",
            label="Contract number",
            value="441-581-246458/25",
            confidence=0.95,
            method="regex",
            source_text="Contract № 441-581-246458/25",
            warnings=[],
        ),
        contract_date=missing_field("contract_date", "Contract date"),
        period_start=missing_field("period_start", "Period start"),
        period_end=missing_field("period_end", "Period end"),
        insurer_inn=missing_field("insurer_inn", "Insurer INN"),
        policyholder_inn=missing_field("policyholder_inn", "Policyholder INN"),
        total_insurance_premium=missing_field(
            "total_insurance_premium",
            "Total insurance premium",
        ),
        total_insured_amount=missing_field(
            "total_insured_amount",
            "Total insured amount",
        ),
        pledge_agreements=[
            ExtractedField(
                name="pledge_agreement",
                label="Pledge agreement",
                value="Agreement 1",
                confidence=1.0,
                method="manual",
                source_text=None,
                warnings=[],
            )
        ],
    )

    result_dict = result.to_dict()

    assert result_dict["contract_number"]["value"] == "441-581-246458/25"
    assert result_dict["contract_date"]["method"] == "not_found"
    assert result_dict["contract_date"]["confidence"] == 0.0
    assert result_dict["pledge_agreements"][0]["method"] == "manual"


def test_parse_docx_keeps_paragraph_and_table_order():
    file_path = Path(__file__).resolve().parent / "ordered_test.docx"
    docx = DocxDocument()
    docx.add_paragraph("First paragraph")

    table = docx.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Policy"
    table.cell(1, 1).text = "123"

    docx.add_paragraph("Last paragraph")
    try:
        docx.save(file_path)

        parsed_document = parse_docx(file_path)

        assert parsed_document.paragraphs == ["First paragraph", "Last paragraph"]
        assert parsed_document.tables == [["Name\tValue", "Policy\t123"]]
        assert parsed_document.raw_blocks == [
            "First paragraph",
            "Name\tValue\nPolicy\t123",
            "Last paragraph",
        ]
        assert parsed_document.full_text == (
            "First paragraph\nName\tValue\nPolicy\t123\nLast paragraph"
        )
    finally:
        file_path.unlink(missing_ok=True)


def test_parse_pdf_extracts_pages_and_word_coordinates(monkeypatch):
    fake_pages = [
        FakePdfPage(
            "First PDF page",
            [
                {
                    "text": "First",
                    "x0": 10,
                    "top": 20,
                    "x1": 40,
                    "bottom": 30,
                }
            ],
        ),
        FakePdfPage("Second PDF page", []),
    ]

    def fake_open(file_path):
        assert file_path == "contract.pdf"
        return FakePdf(fake_pages)

    monkeypatch.setattr(pdf_parser.pdfplumber, "open", fake_open)

    parsed_document = parse_pdf("contract.pdf")

    assert parsed_document.full_text == "First PDF page\n\nSecond PDF page"
    assert parsed_document.warning is None
    assert len(parsed_document.pages) == 2
    assert parsed_document.pages[0].page_number == 1
    assert parsed_document.pages[0].words[0].text == "First"
    assert parsed_document.pages[0].words[0].x0 == 10.0
    assert parsed_document.pages[0].words[0].page_number == 1
    assert parsed_document.pages[1].page_number == 2


def test_parse_pdf_returns_warning_for_scanned_pdf(monkeypatch):
    def fake_open(file_path):
        return FakePdf([FakePdfPage(None, [])])

    monkeypatch.setattr(pdf_parser.pdfplumber, "open", fake_open)

    parsed_document = parse_pdf("scan.pdf")

    assert parsed_document.full_text == ""
    assert parsed_document.pages[0].text == ""
    assert parsed_document.pages[0].words == []
    assert parsed_document.warning == SCAN_WARNING
