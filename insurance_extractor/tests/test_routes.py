import json
from io import BytesIO
from types import SimpleNamespace

from docx import Document as DocxDocument

from app import routes
from app.models import Document


def make_docx_file() -> BytesIO:
    docx = DocxDocument()
    docx.add_paragraph("Договор страхования")
    docx.add_paragraph("Страхователь: Иванов Иван")

    file_stream = BytesIO()
    docx.save(file_stream)
    file_stream.seek(0)
    return file_stream


def make_extractable_docx_file() -> BytesIO:
    docx = DocxDocument()
    docx.add_paragraph("ДОГОВОР СТРАХОВАНИЯ ИМУЩЕСТВА № 78542/919/50117/25")
    docx.add_paragraph("г. Волгоград «13» марта 2026 г.")
    docx.add_paragraph('Страховщик: АО "Страховая компания", ИНН 7701234567')
    docx.add_paragraph('Страхователь: ООО "Ромашка", ИНН 771234567890')
    docx.add_paragraph("ПЕРИОД СТРАХОВАНИЯ")
    docx.add_paragraph("с «30» ноября 2025 года по «29» ноября 2026 года")
    docx.add_paragraph("Итого общая страховая сумма по Договору: 73 353 000,00")
    docx.add_paragraph("Итого общая страховая премия по настоящему Договору: 73 353,00")
    docx.add_paragraph("Имущество является предметом залога по Договору залога № СВЛ/111111-111111-303")

    file_stream = BytesIO()
    docx.save(file_stream)
    file_stream.seek(0)
    return file_stream


def test_index_page_opens(client):
    response = client.get("/")

    assert response.status_code == 200
    assert "Договоры страхования".encode("utf-8") in response.data


def test_docx_upload_creates_document(client, app):
    response = client.post(
        "/",
        data={"document": (make_docx_file(), "contract.docx")},
        content_type="multipart/form-data",
        follow_redirects=True,
    )

    assert response.status_code == 200
    assert "Договор страхования".encode("utf-8") in response.data

    with app.app_context():
        document = Document.query.one()
        assert document.original_filename == "contract.docx"
        assert "Страхователь" in document.text_content


def test_pdf_upload_creates_document(client, app, monkeypatch):
    def fake_parse_pdf(file_path):
        return SimpleNamespace(full_text="PDF contract text", warning=None)

    monkeypatch.setattr(routes, "parse_pdf", fake_parse_pdf)

    response = client.post(
        "/",
        data={"document": (BytesIO(b"fake pdf content"), "contract.pdf")},
        content_type="multipart/form-data",
        follow_redirects=True,
    )

    assert response.status_code == 200
    assert b"PDF contract text" in response.data

    with app.app_context():
        document = Document.query.one()
        assert document.original_filename == "contract.pdf"
        assert document.text_content == "PDF contract text"


def test_docx_upload_stores_extraction_result_and_renders_fields(client, app):
    response = client.post(
        "/",
        data={"document": (make_extractable_docx_file(), "extractable.docx")},
        content_type="multipart/form-data",
        follow_redirects=True,
    )

    assert response.status_code == 200
    assert b"78542/919/50117/25" in response.data
    assert "Скопировать".encode("utf-8") in response.data
    assert "Скопировать всё через Tab".encode("utf-8") in response.data

    with app.app_context():
        document = Document.query.one()
        extraction_data = json.loads(document.extraction_result_json)
        assert extraction_data["contract_number"]["value"] == "78542/919/50117/25"
        assert extraction_data["contract_date"]["value"] == "13.03.2026"
        assert extraction_data["period_start"]["value"] == "30.11.2025"
        assert extraction_data["period_end"]["value"] == "29.11.2026"
        assert extraction_data["total_insured_amount"]["value"] == "73353000.00"
        assert extraction_data["total_insurance_premium"]["value"] == "73353.00"
        assert extraction_data["pledge_agreements"][0]["value"] == "СВЛ/111111-111111-303"
